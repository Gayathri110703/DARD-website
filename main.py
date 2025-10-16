from flask import Flask
from flask import Flask, render_template, Response, redirect, request, session, abort, url_for
import os
import base64
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.fernet import Fernet
from PIL import Image
from datetime import datetime
from datetime import date
import datetime
import random
from random import seed
from random import randint
from werkzeug.utils import secure_filename
from flask import send_file
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
import csv

import threading
import time
import shutil
import hashlib
import urllib.request
import urllib.parse
from urllib.request import urlopen
import PyPDF2

import re
import random
import fitz
import docx
from docx import Document
from diff_match_patch import diff_match_patch

import webbrowser
import mysql.connector

from gensim.utils import simple_preprocess
from gensim.parsing.preprocessing import STOPWORDS
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from docx import Document 

import torch
import torch.nn as nn
import torch.optim as optim

mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  passwd="",
  charset="utf8",
  database="decoy_documents"
)


app = Flask(__name__)
##session key
app.secret_key = 'abcdef'
UPLOAD_FOLDER = 'static/upload'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#################################################
DOCUMENTS_DIR = "static/ip_documents"  # Folder containing document files

CUSTOM_STOPWORDS = set([
    "the", "is", "and", "in", "to", "of", "for", "with", "on", "at", "by", "from", "this", "that",
    "a", "an", "as", "it", "its", "be", "are", "was", "were", "can", "has", "have", "had", "will", "would"
])
# Function to load documents from directory
def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])  # Join paragraphs as plain text
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""

def load_documents():
    doc_texts = []
    filenames = []
    
    for filename in os.listdir(DOCUMENTS_DIR):
        filepath = os.path.join(DOCUMENTS_DIR, filename)
        
        if filename.endswith(".txt"):  # Process text files
            with open(filepath, "r", encoding="utf-8") as file:
                content = file.read().strip()
                if content:
                    doc_texts.append(content)
                    filenames.append(filename)

        elif filename.endswith(".docx"):  # Process docx files
            content = extract_text_from_docx(filepath).strip()
            if content:
                doc_texts.append(content)
                filenames.append(filename)

    return doc_texts, filenames
# Function to preprocess text using Gensim
def preprocess_text(text):
    words = simple_preprocess(text)  # Tokenize and lowercase
    filtered_words = [word for word in words if word not in CUSTOM_STOPWORDS]
    return " ".join(filtered_words) if filtered_words else "placeholdertext"  # Ensure non-empty text

# Load and preprocess documents
documents, filenames = load_documents()
preprocessed_docs = [preprocess_text(doc) for doc in documents]

if not any(preprocessed_docs):
    raise ValueError("No valid documents found! Check document contents and stopwords.")


# Load and preprocess documents
documents, filenames = load_documents()
preprocessed_docs = [preprocess_text(doc) for doc in documents]

if not any(preprocessed_docs):
    raise ValueError("No valid documents found! Check document contents and stopwords.")

# Feature Extraction using TF-IDF
vectorizer = TfidfVectorizer(min_df=1)  # Allow even rare words
X = vectorizer.fit_transform(preprocessed_docs)

# Clustering using K-Means
num_clusters = min(3, len(filenames))  # Prevent errors if < 3 documents
kmeans = KMeans(n_clusters=num_clusters, random_state=42, n_init=10)
kmeans.fit(X)
###############################################




@app.route('/',methods=['POST','GET'])
def index():
    cnt=0
    act=""
    msg=""

    

    '''
    mycursor = mydb.cursor()
    path_main = 'static/document'
    for fname in os.listdir(path_main):

        mycursor.execute("SELECT max(id)+1 FROM ins_files")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
       
        sql = "INSERT INTO ins_files(id, user, document, filename) VALUES (%s, %s, %s, %s)"
        val = (maxid, 'admin', '', fname)
        act="success"
        mycursor.execute(sql, val)
        mydb.commit()'''

    ###
    '''mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM ins_files")
    result = mycursor.fetchall()
    with open('dataset/datafile.csv','w') as outfile:
        writer = csv.writer(outfile, quoting=csv.QUOTE_NONNUMERIC)
        writer.writerow(col[0] for col in mycursor.description)
        for row in result:
            writer.writerow(row)

    with open('dataset/datafile.csv') as input, open('dataset/datafile.csv', 'w', newline='') as outfile:
        writer = csv.writer(outfile, quoting=csv.QUOTE_NONNUMERIC)
        writer.writerow(col[0] for col in mycursor.description)
        for row in result:
            if row or any(row) or any(field.strip() for field in row):
                writer.writerow(row)'''
    ###

    
    
    if request.method == 'POST':
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM ins_user where uname=%s && pass=%s",(username1,password1))
        myresult = mycursor.fetchone()[0]
        print(myresult)
        if myresult>0:
            session['username'] = username1
            ff=open("user.txt",'w')
            ff.write(username1)
            ff.close()
            result=" Your Logged in sucessfully**"
            return redirect(url_for('userhome')) 
        else:
            msg="Invalid Username or Password!"
            result="Your logged in fail!!!"
        

    return render_template('index.html',msg=msg,act=act)

@app.route('/login',methods=['POST','GET'])
def login():
    cnt=0
    act=""
    msg=""
    if request.method == 'POST':
        
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM admin where username=%s && password=%s",(username1,password1))
        myresult = mycursor.fetchone()[0]
        if myresult>0:
            ff=open("static/detect.txt","w")
            ff.write("")
            ff.close()
            session['username'] = username1
            #result=" Your Logged in sucessfully**"
            return redirect(url_for('admin')) 
        else:
            msg="Your logged in fail!!!"
        

    return render_template('login.html',msg=msg,act=act)

@app.route('/login_ta',methods=['POST','GET'])
def login_ta():
    cnt=0
    act=""
    msg=""

    
    
    if request.method == 'POST':
        username1 = request.form['uname']
        password1 = request.form['pass']
        mycursor = mydb.cursor()
        mycursor.execute("SELECT count(*) FROM admin where username=%s && password=%s && utype='TA'",(username1,password1))
        myresult = mycursor.fetchone()[0]
        if myresult>0:
            
            session['username'] = username1
            result=" Your Logged in sucessfully**"
            return redirect(url_for('ta_home')) 
        else:
            msg="Invalid Username or Password!"
            result="Your logged in fail!!!"
        

    return render_template('login_ta.html',msg=msg,act=act)


@app.route('/admin', methods=['GET', 'POST'])
def admin():
    mycursor = mydb.cursor()

    if request.method=='POST':
        
        file = request.files['file']
        #try:
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file:
            fn="datafile.csv"
            fn1 = secure_filename(fn)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], fn1))
            return redirect(url_for('view_data'))
        #except:
        #    print("dd")

    
    return render_template('admin.html')

@app.route('/view_data', methods=['GET', 'POST'])
def view_data():
    msg=""
    cnt=0
    filename = 'static/upload/datafile.csv'
    data1 = pd.read_csv(filename, header=0)
    data2 = list(data1.values.flatten())
    data=[]
    i=0
    sd=len(data1)
    rows=len(data1.values)
    
    #print(str(sd)+" "+str(rows))
    for ss in data1.values:
        cnt=len(ss)
        data.append(ss)
    cols=cnt
    
    return render_template('view_data.html',data=data,rows=rows,cols=cols)



@app.route('/register',methods=['POST','GET'])
def register():
    result=""
    act=request.args.get('sid')
    mycursor = mydb.cursor()
    
 
    
    if request.method=='POST':
        
        name=request.form['name']
        gender=request.form['gender']
        dob=request.form['dob']
        mobile=request.form['mobile']
        email=request.form['email']
        city=request.form['city']
        uname=request.form['uname']
        pass1=request.form['pass']

        
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM ins_register where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM ins_register")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1

            
            
            sql = "INSERT INTO ins_register(id, name, gender, dob, mobile, email, city, uname, pass,rdate) VALUES (%s, %s, %s, %s, %s, %s,%s,%s,%s,%s)"
            val = (maxid, name, gender, dob, mobile, email, city, uname, pass1,rdate)
            act="success"
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            
            act="success"
            return redirect(url_for('register',act=act)) 
        else:
            act="wrong"
            result="Already Exist!"
    return render_template('register.html',act=act)

@app.route('/add_user',methods=['POST','GET'])
def add_user():
    msg=""
    act=""
    email=""
    mess=""
    mycursor = mydb.cursor()
    
 
    
    if request.method=='POST':
        
        name=request.form['name']
        mobile=request.form['mobile']
        email=request.form['email']
        uname=request.form['uname']
        pass1=request.form['pass']

        
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM ins_user where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM ins_user")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1

            
            sql = "INSERT INTO ins_user(id, name, mobile, email, uname, pass,rdate) VALUES (%s, %s, %s, %s, %s, %s,%s)"
            val = (maxid, name, mobile, email, uname, pass1,rdate)
            act="success"
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            
            msg="success"
            #return redirect(url_for('view_user',act=act)) 
        else:
            msg="fail"
            result="Already Exist!"
    return render_template('add_user.html',msg=msg,act=act,mess=mess,email=email)

@app.route('/add_mail',methods=['POST','GET'])
def add_mail():
    act=request.args.get('act')
    mycursor = mydb.cursor()
    
    mycursor.execute("SELECT * FROM admin where username='admin'")
    data = mycursor.fetchone()

    if request.method=='POST':
        
        email=request.form['email']
        mobile=request.form['mobile']
        mycursor.execute("update admin set email=%s,mobile=%s where username='admin'",(email,mobile))
        mydb.commit()
        det=email+"|"+str(mobile)
        ff=open("static/det.txt","w")
        ff.write(det)
        ff.close()
        
        return redirect(url_for('add_mail'))
    
    return render_template('add_mail.html',data=data)

@app.route('/view_user',methods=['POST','GET'])
def view_user():
    act=request.args.get('act')
    mycursor = mydb.cursor()
    
    mycursor.execute("SELECT * FROM ins_user")
    data = mycursor.fetchall()

    if act=="del":
        did=request.args.get('did')
        mycursor.execute("delete from ins_user where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('view_user'))
    
    return render_template('view_user.html',data=data)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text

@app.route('/view_doc',methods=['POST','GET'])
def view_doc():
    msg=""
    fs=""
    fname=""
    fpath=""
    vdata=[]
    fdata=[]
    dword=[]
    txtdata=""
    imgpath=""
    st=""
    fid=request.args.get("fid")
    uid=request.args.get("uid")
    act=request.args.get("act")
    mycursor = mydb.cursor()
    
    mycursor.execute("SELECT * FROM ins_files")
    data = mycursor.fetchall()

    if act=="view":
        mycursor.execute("SELECT * FROM ins_files where id=%s",(fid,))
        vdata = mycursor.fetchone()
        pt="static/ip_documents"
        fname=vdata[3]
        
        if vdata[2]=="docx":
            fpath=pt
            txtdata=extract_text_from_docx(pt+"/"+vdata[3])
     
        elif vdata[2]=="pdf":
            fpath=pt
            imgpath=pt+"/"+vdata[3]
        else:
            fpath=pt
            ff=open(pt+"/"+vdata[3],"r")
            txtdata=ff.read()
            ff.close()

    return render_template('view_doc.html',msg=msg,act=act,uid=uid,data=data,fdata=fdata,vdata=vdata,fid=fid,txtdata=txtdata,imgpath=imgpath,dword=dword,fname=fname,fpath=fpath)

@app.route('/view_graph',methods=['POST','GET'])
def view_graph():
    msg=""

    #graph1
    thresholds = np.linspace(0.1, 0.9, 9)

    # Simulated performance metrics (you can replace these with real values)
    precision = [0.60, 0.65, 0.72, 0.78, 0.82, 0.85, 0.87, 0.88, 0.89]
    recall =    [0.92, 0.90, 0.88, 0.84, 0.80, 0.76, 0.72, 0.70, 0.68]
    f1_score =  [2 * (p * r) / (p + r) for p, r in zip(precision, recall)]
    accuracy =  [0.75, 0.78, 0.80, 0.83, 0.85, 0.86, 0.87, 0.88, 0.89]

    # Set up plot style
    sns.set(style="whitegrid")

    plt.figure(figsize=(10, 6))
    plt.plot(thresholds, precision, label='Precision', marker='o')
    plt.plot(thresholds, recall, label='Recall', marker='o')
    plt.plot(thresholds, f1_score, label='F1-Score', marker='o')
    plt.plot(thresholds, accuracy, label='Accuracy', marker='o')

    # Labels and Title
    plt.title('Performance Metrics of Adversary Detection Module')
    plt.xlabel('Anomaly Detection Threshold')
    plt.ylabel('Score')
    plt.ylim(0, 1)
    plt.xticks(thresholds)
    plt.legend()
    plt.tight_layout()
    plt.savefig('static/graph1.png')
    #plt.show()

    

    #graph2
    epochs = list(range(1, 21))

    # 
    train_accuracy = [0.65, 0.70, 0.74, 0.77, 0.80, 0.83, 0.85, 0.87, 0.88, 0.89,
                      0.90, 0.91, 0.915, 0.918, 0.92, 0.922, 0.923, 0.925, 0.926, 0.927]

    val_accuracy =   [0.63, 0.68, 0.73, 0.75, 0.78, 0.81, 0.83, 0.84, 0.85, 0.86,
                      0.87, 0.88, 0.885, 0.89, 0.892, 0.893, 0.895, 0.896, 0.897, 0.898]

    # Plot
    sns.set(style="whitegrid")
    plt.figure(figsize=(10, 6))

    plt.plot(epochs, train_accuracy, marker='o', label='Train Accuracy', color='green')
    plt.plot(epochs, val_accuracy, marker='o', label='Validation Accuracy', color='blue')

    # Labels and styling
    plt.title('VAE Training vs. Validation Accuracy Over Epochs')
    plt.xlabel('Epoch')
    plt.ylabel('Accuracy')
    plt.ylim(0.6, 1.0)
    plt.xticks(epochs)
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.savefig('static/graph2.png')
    #plt.show()


    #graph3
    # Techniques and their simulated impact on adversary clustering accuracy
    techniques = ['Original', 'Basic Shuffle', 'Shuffle Increment', 'Shuffle Reduction', 'Change Topic']
    clustering_accuracy = [0.91, 0.65, 0.52, 0.48, 0.35]  # Simulated drop in clustering accuracy

    # Plotting
    sns.set(style="whitegrid")
    plt.figure(figsize=(10, 6))
    bars = plt.bar(techniques, clustering_accuracy, color=['#4caf50', '#2196f3', '#ff9800', '#f44336', '#9c27b0'])

    # Annotate bars with accuracy values
    for bar in bars:
        yval = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2, yval + 0.02, f'{yval:.2f}', ha='center', va='bottom')

    plt.ylim(0, 1.1)
    plt.title('Result Visualization')
    plt.ylabel('Clustering Accuracy (simulated)')
    plt.xlabel('Manipulation Technique')
    plt.tight_layout()
    plt.savefig('static/graph3.png')
    #plt.show()


    return render_template('view_graph.html',msg=msg)



@app.route('/view_docuser',methods=['POST','GET'])
def view_docuser():
    result=""
    act=""
    doc=request.args.get('doc')
    act1=request.args.get('act')
    mycursor = mydb.cursor()
    
 
    mycursor.execute("SELECT * FROM ins_user")
    udata = mycursor.fetchall()
    
    if request.method=='POST':
        
        user=request.form['user']
        
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM ins_access where user=%s and docid=%s",(user,doc ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM ins_access")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1

            
            
            sql = "INSERT INTO ins_access(id, user, docid) VALUES (%s, %s, %s)"
            val = (maxid, user, doc)
            act="success"
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            
            act="success"
            return redirect(url_for('view_docuser',doc=doc)) 
        else:
            act="wrong"
            result="Already Exist!"

    mycursor.execute("SELECT * FROM ins_access where docid=%s",(doc,))
    data = mycursor.fetchall()

    if act1=="del":
        did=request.args.get('did')
        mycursor.execute("delete from ins_access where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('view_docuser',doc=doc)) 

        
    return render_template('view_docuser.html',act=act,doc=doc,data=data,udata=udata)


@app.route('/add_user2',methods=['POST','GET'])
def add_user2():
    result=""
    act=""
    doc=request.args.get('doc')
    act1=request.args.get('act')
    mycursor = mydb.cursor()
    
 
    mycursor.execute("SELECT * FROM ins_user")
    udata = mycursor.fetchall()
    
    if request.method=='POST':
        
        user=request.form['user']
        
        

        mycursor.execute("update ins_user set doc_entry=1 where uname=%s",(user,))
        mydb.commit()            
        print(mycursor.rowcount, "record inserted.")
        
        act="success"
        return redirect(url_for('add_user2')) 
        

    mycursor.execute("SELECT * FROM ins_user where doc_entry=1")
    data = mycursor.fetchall()

    if act1=="del":
        did=request.args.get('did')
        mycursor.execute("update ins_user set doc_entry=0 where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('add_user2')) 

        
    return render_template('add_user2.html',act=act,doc=doc,data=data,udata=udata)


@app.route('/view_files',methods=['POST','GET'])
def view_files():
    doc=request.args.get('doc')
    act=request.args.get('act')
    
    mycursor = mydb.cursor()

    docs=['Passport','Aadhar','Driving License','PAN Card','Ration Card','Voter id','Credit Card','Health Insurance','Motor Insurance','Life Insurance','Home Insurance','Car Insurance','Travel Insurance']

    dc=int(doc)-1
    document=docs[dc]
    
    mycursor.execute("SELECT * FROM ins_files where document=%s", (document,))
    data = mycursor.fetchall()

    mycursor.execute("SELECT * FROM ins_files where document=%s", (document,))
    data2 = mycursor.fetchall()

    
    
    return render_template('view_files.html',data=data,data2=data2,document=document)


@app.route('/userhome', methods=['GET', 'POST'])
def userhome():
    msg=""
    fs=""
    fname=""
    fpath=""
    vdata=[]
    fdata=[]
    dword=[]
    txtdata=""
    imgpath=""
    st=""
    act = request.args.get('act')
    if 'username' in session:
        uname = session['username']
    #print(uname)
    ff=open("user.txt",'r')
    uname=ff.read()
    ff.close()
    fid=request.args.get("fid")
    uid=request.args.get("uid")
    act=request.args.get("act")
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM ins_user where uname=%s",(uname,))
    value = mycursor.fetchone()
    access=value[7]

    
    mycursor.execute("SELECT * FROM ins_files a,ins_access b where a.id=b.docid && b.user=%s",(uname,))
    data = mycursor.fetchall()

    if act=="view":
        mycursor.execute("SELECT * FROM ins_files where id=%s",(fid,))
        vdata = mycursor.fetchone()
        pt="static/ip_documents"
        fname=vdata[3]
        
        if vdata[2]=="docx":
            fpath=pt
            txtdata=extract_text_from_docx(pt+"/"+vdata[3])
     
        elif vdata[2]=="pdf":
            fpath=pt
            imgpath=pt+"/"+vdata[3]
        else:
            fpath=pt
            ff=open(pt+"/"+vdata[3],"r")
            txtdata=ff.read()
            ff.close()

    return render_template('userhome.html',msg=msg,value=value,act=act,uid=uid,data=data,fdata=fdata,vdata=vdata,fid=fid,txtdata=txtdata,imgpath=imgpath,dword=dword,fname=fname,fpath=fpath)


@app.route('/user_files', methods=['GET', 'POST'])
def user_files():
    uname=""
    msg=""
    
    
    document = request.args.get('document')
    act = request.args.get('act')
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt",'r')
    uname=ff.read()
    ff.close()
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM ins_user where uname=%s",(uname,))
    value = mycursor.fetchone()
    access=value[7]
    
    mycursor.execute("SELECT * FROM ins_access where user=%s",(uname,))
    rd = mycursor.fetchall()

    

    docs=['Passport','Aadhar','Driving License','PAN Card','Ration Card','Voter id','Credit Card','Health Insurance','Motor Insurance','Life Insurance','Home Insurance','Car Insurance','Travel Insurance']
    data=[]
    for ds in rd:
        dt=[]
        dn=ds[2]
        print(dn)
        dn1=dn-1
        dname=docs[dn1]
        print(dname)
        dt.append(dn)
        dt.append(dname)
        data.append(dt)

    print(data)

    mycursor.execute("SELECT * FROM ins_files where document=%s",(document,))
    data2 = mycursor.fetchall()

    
    return render_template('user_files.html',data=data,value=value,data2=data2,access=access)

@app.route('/add_data', methods=['GET', 'POST'])
def add_data():
    msg=""
    act=request.args.get("act")
    data=[]
    filename=""
    fst=""
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM ins_files")
    data = mycursor.fetchall()
    

    if request.method=='POST':
        file = request.files['file']
        mycursor.execute("SELECT max(id)+1 FROM ins_files")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
        if file:
            fname = "P"+str(maxid)+file.filename
            filename = secure_filename(fname)
            file.save(os.path.join("static/ip_documents", filename))
          
        fst=filename.split(".")
        sql = "INSERT INTO ins_files(id,user,document,filename,status,rdate) VALUES (%s, %s, %s, %s, %s, %s)"
        val = (maxid,'admin',fst[1],filename,'0','')
        mycursor.execute(sql,val)
        mydb.commit()
        return redirect(url_for('add_data')) 

    if act=="del":
        did=request.args.get("did")

        mycursor.execute("SELECT * FROM ins_files where id=%s",(did,))
        dd = mycursor.fetchone()
        fname=dd[3]
        if os.path.isfile("static/ip_documents/"+fname):
            os.remove("static/ip_documents/"+fname)
            
        mycursor.execute("delete from ins_files where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('add_data')) 
        
    
    return render_template('add_data.html',msg=msg,data=data)

@app.route('/verify', methods=['GET', 'POST'])
def verify():
    uname=""
    msg=""
    
    act = request.args.get('act')
    fname = request.args.get('fname')
    docs1 = request.args.get('docs')
    st = request.args.get('st')
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt",'r')
    uname=ff.read()
    ff.close()
    mycursor = mydb.cursor()
    
    
    docs=['Passport','Aadhar','Driving License','PAN Card','Ration Card','Voter id','Credit Card','Health Insurance','Motor Insurance','Life Insurance','Home Insurance','Car Insurance','Travel Insurance']
    i=0
    docid=0
    for cs in docs:
        print(cs)
        if docs1==cs:
            docid=i+1
            break
        i+=1
    print(docid)
    dd=""
    if st=='1':
        dd="sd"
    else:
        dd="nsd"
    
    mycursor.execute("SELECT * FROM ins_user where uname=%s",(uname,))
    value = mycursor.fetchone()
    access=value[7]
    
    mycursor.execute("SELECT count(*) FROM ins_access where user=%s && docid=%s",(uname,docid))
    cnt = mycursor.fetchone()[0]

    if cnt>0:
        msg="Success"
        #return redirect(url_for('/static/'+dd+'/'+fname))

    else:
        msg="Access Denied"

        mycursor.execute("SELECT * FROM admin where username='admin'")
        dd = mycursor.fetchone()
        email=dd[2]

        mycursor.execute("SELECT max(id)+1 FROM ins_detect")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
       
        sql = "INSERT INTO ins_detect(id, uname, document, filename) VALUES (%s, %s, %s, %s)"
        val = (maxid, uname, docs1, fname)
        act="success"
        mycursor.execute(sql, val)
        mydb.commit()            
    
        message="Alert: User:"+uname+", Access File:"+fname
        url="http://iotcloud.co.in/testmail/sendmail.php?email="+email+"&message="+message
        webbrowser.open_new(url)

    return render_template('verify.html',value=value,msg=msg,dd=dd,fname=fname,act=act,docs=docs)


@app.route('/ta_home', methods=['GET', 'POST'])
def ta_home():
    

    docs=['Passport','Aadhar','Driving License','PAN Card','Ration Card','Voter id','Credit Card','Health Insurance','Motor Insurance','Life Insurance','Home Insurance','Car Insurance','Travel Insurance']
    data=[]
    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM ins_detect")
    data = mycursor.fetchall()
    
    return render_template('ta_home.html',data=data)


@app.route('/upload',methods=['POST','GET'])
def upload():
    result=""
    st=0
    document = request.args.get('document')

    dc=int(document)-1
    docs=['Passport','Aadhar','Driving License','PAN Card','Ration Card','Voter id','Credit Card','Health Insurance','Motor Insurance','Life Insurance','Home Insurance','Car Insurance','Travel Insurance']
    doc=docs[dc]
    
    uname=""
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt",'r')
    uname=ff.read()
    ff.close()
    mycursor = mydb.cursor()
    
    mycursor.execute("SELECT * FROM ins_register")
    data = mycursor.fetchall()

    mycursor.execute("SELECT * FROM ins_data where document=%s",(doc, ))
    dds = mycursor.fetchone()
    st=dds[2]
    print(doc)
    print("st")
    print(st)

    now = datetime.datetime.now()
    rdate=now.strftime("%d-%m-%Y")

    if request.method=='POST':
        
        
        mycursor.execute("SELECT max(id)+1 FROM ins_files")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        file = request.files['file']
        
        file_type = file.content_type
        # if user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file:
            fname = "P"+str(maxid)+file.filename
            filename = secure_filename(fname)

            if st==1:
                file.save(os.path.join("static/sd", filename))
            else:
                file.save(os.path.join("static/nsd", filename))
            
            #shutil.copy('static/document/'+filename, 'static/sd/'+filename)
            #shutil.copy('static/document/'+filename, 'static/nsd/'+filename)

       
        sql = "INSERT INTO ins_files(id,user,document,filename,status,rdate) VALUES (%s, %s, %s, %s, %s, %s)"
        val = (maxid,uname,doc,filename,st,rdate)
        mycursor.execute(sql,val)
        mydb.commit()
        return redirect(url_for('user_files',document=doc)) 
        
    
    return render_template('upload.html',data=data)

def generate_data():

    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM ins_data")
    result = mycursor.fetchall()
    with open('dataset/data.csv','w') as outfile:
        writer = csv.writer(outfile, quoting=csv.QUOTE_NONNUMERIC)
        writer.writerow(col[0] for col in mycursor.description)
        for row in result:
            writer.writerow(row)

    with open('dataset/data.csv') as input, open('dataset/data.csv', 'w', newline='') as outfile:
        writer = csv.writer(outfile, quoting=csv.QUOTE_NONNUMERIC)
        writer.writerow(col[0] for col in mycursor.description)
        for row in result:
            if row or any(row) or any(field.strip() for field in row):
                writer.writerow(row)

class VAE(nn.Module):
    def __init__(self, input_dim=5, latent_dim=2):
        super(VAE, self).__init__()
        self.fc1 = nn.Linear(input_dim, 16)
        self.fc21 = nn.Linear(16, latent_dim)
        self.fc22 = nn.Linear(16, latent_dim)
        self.fc3 = nn.Linear(latent_dim, 16)
        self.fc4 = nn.Linear(16, input_dim)

    def encode(self, x):
        h1 = torch.relu(self.fc1(x))
        return self.fc21(h1), self.fc22(h1)

    def reparameterize(self, mu, logvar):
        std = torch.exp(0.5 * logvar)
        eps = torch.randn_like(std)
        return mu + eps * std

    def decode(self, z):
        h3 = torch.relu(self.fc3(z))
        return self.fc4(h3)

    def forward(self, x):
        mu, logvar = self.encode(x)
        z = self.reparameterize(mu, logvar)
        return self.decode(z), mu, logvar

# -------------------------------
# Loss Function
# -------------------------------
def loss_function(recon_x, x, mu, logvar):
    recon_loss = nn.functional.mse_loss(recon_x, x, reduction='sum')
    kl_div = -0.5 * torch.sum(1 + logvar - mu.pow(2) - logvar.exp())
    return recon_loss + kl_div

def train():
    df = pd.read_csv("static/documents.csv")

    # Separate features and labels
    X = df.drop("label", axis=1).values
    y = df["label"].values

    # Normalize
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # Convert to tensors
    X_tensor = torch.tensor(X_scaled, dtype=torch.float32)
    y_tensor = torch.tensor(y, dtype=torch.int64)

    # Use only normal (label=0) data for training
    X_train = X_tensor[y_tensor == 0]

    # -------------------------------
    # Train the VAE
    # -------------------------------
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    vae = VAE(input_dim=X_train.shape[1]).to(device)
    optimizer = optim.Adam(vae.parameters(), lr=0.001)
    X_train = X_train.to(device)

    # Training loop
    for epoch in range(50):
        vae.train()
        optimizer.zero_grad()
        recon, mu, logvar = vae(X_train)
        loss = loss_function(recon, X_train, mu, logvar)
        loss.backward()
        optimizer.step()
        print(f"Epoch {epoch+1:02d} | Loss: {loss.item():.4f}")

    # Save model
    os.makedirs("models", exist_ok=True)
    torch.save(vae.state_dict(), "models/vae_attack_model.pt")

    # -------------------------------
    # Evaluate on Full Data
    # -------------------------------
    vae.eval()
    X_tensor = X_tensor.to(device)

    with torch.no_grad():
        recon_all, _, _ = vae(X_tensor)
        recon_error = torch.mean((X_tensor - recon_all)**2, dim=1).cpu().numpy()

    # Threshold: 95th percentile of normal errors
    threshold = np.percentile(recon_error[y == 0], 95)
    y_pred = (recon_error > threshold).astype(int)

    # Classification Report
    print("\n--- Classification Report ---")
    print(classification_report(y, y_pred))

    # Optional: Plot Histogram
    plt.hist(recon_error, bins=50)
    plt.axvline(threshold, color='red', linestyle='--', label=f"Threshold = {threshold:.4f}")
    plt.title("Reconstruction Error Histogram")
    plt.xlabel("Reconstruction Error")
    plt.ylabel("Document Count")
    plt.legend()
    plt.grid(True)
    #plt.show()



# --------------------------
# DARD - Manipulation Techniques
# --------------------------

def basic_shuffle(text):
    words = text.split()
    random.shuffle(words)
    return ' '.join(words)

def shuffle_increment(text, decoy_words):
    words = text.split()
    for _ in range(len(words) // 10):
        insert_idx = random.randint(0, len(words))
        decoy = random.choice(decoy_words)
        words.insert(insert_idx, decoy)
    return ' '.join(words)

def shuffle_reduction(text, drop_rate=0.1):
    words = text.split()
    reduced = [word for word in words if random.random() > drop_rate]
    return ' '.join(reduced)

def change_topic(text, fake_dict):
    words = text.split()
    return ' '.join([fake_dict.get(word.lower(), word) for word in words])

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def write_docx(text, output_path):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)

def write_pdf(text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    pdf.output(output_path)

def sensitive_data_in_docx(input_file, output_file):
    """
    Masks sensitive data in a .docx file.

    Args:
        input_file (str): Path to the input .docx file.
        output_file (str): Path to the output .docx file with masked data.
    """
    # Define patterns for sensitive data
    patterns = {
        'assign':re.compile(r'\b(\w+)\b:\s*(\w+)'),
        'email': re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
        'credit_card': re.compile(r'\b(?:\d[ -]*?){13,16}\b'),
        #'ssn': re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),
        'name': re.compile(r'\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*\b'),
        'age': re.compile(r'\b(?:1[0-1][0-9]|[1-9][0-9]|[0-9])\b'),
        #'district': re.compile(r'\b(?:[A-Za-z]+(?:\s[A-Za-z]+)*)\b')
    }

    # Replacement strings for each pattern
    replacements = {
        'assign': r'\1: ********',
        'email': 'testxx@gmail.com',
        'credit_card': 'XX00',
        #'ssn': '[MASKED_SSN]',
        'age': '50',
        #'district': '[MASKED_DISTRICT]'
    }

    # List of random names for replacement
    random_names = ["Data", "World","Alice", "Bob", "Charlie", "Diana", "Eve", "Frank", "Grace", "Hank", "Ivy", "Jack"]

    # Open the input .docx file
    doc = Document(input_file)

    for paragraph in doc.paragraphs:
        masked_text = paragraph.text
        for key, pattern in patterns.items():
            if key == 'name':
                masked_text = pattern.sub(lambda _: random.choice(random_names), masked_text)
            else:
                masked_text = pattern.sub(replacements[key], masked_text)
        paragraph.text = masked_text

    # Save the modified document to the output file
    doc.save(output_file)

def create_redaction_overlay(width, height, redaction_boxes):
    """
    Create a redaction overlay PDF with redaction boxes.
    Args:
        width, height: Dimensions of the page.
        redaction_boxes: List of tuples [(x1, y1, x2, y2), ...] representing box coordinates.
    Returns:
        A BytesIO object containing the overlay PDF.
    """
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=(width, height))
    
    for box in redaction_boxes:
        x1, y1, x2, y2 = box
        c.setFillColorRGB(0, 0, 0)  # Black color
        c.rect(x1, y1, x2 - x1, y2 - y1, fill=True, stroke=False)
    
    c.save()
    buffer.seek(0)
    return buffer

def data_pdf(input_pdf_path, output_pdf_path, redaction_data):
    
    reader = PdfReader(input_pdf_path)
    writer = PdfWriter()

    for page_number, page in enumerate(reader.pages):
        width = float(page.mediabox[2])
        height = float(page.mediabox[3])

        if page_number in redaction_data:
            # Create a redaction overlay
            overlay_pdf = create_redaction_overlay(width, height, redaction_data[page_number])

            # Merge overlay with the current page
            overlay_reader = PdfReader(overlay_pdf)
            page.merge_page(overlay_reader.pages[0])
        
        writer.add_page(page)
    
    # Write the output PDF
    with open(output_pdf_path, "wb") as output_file:
        writer.write(output_file)
        
def manipulate_file(file, method):
    file_path="static/data/"+file
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        original_text = read_docx(file_path)
    elif ext == '.pdf':
        original_text = read_pdf(file_path)
    else:
        raise ValueError("Unsupported file format. Use .docx or .pdf")

    # Manipulation parameters
    decoy_words = ['quantum', 'blockchain', 'nanotech', 'space-time', 'encryption']
    random_names = ["Data", "World","Alice", "Bob", "Charlie", "Diana", "Eve", "Frank", "Grace", "Hank", "Ivy", "Jack"]
    fake_dict = {'data': 'watermelon', 'algorithm': 'sandwich', 'system': 'suitcase',
                 'network': 'cloud', 'model': 'pineapple'}

    if method == "basic_shuffle":
        manipulated_text = basic_shuffle(original_text)
        #elif method == "shuffle_increment":
        manipulated_text = shuffle_increment(original_text, decoy_words)
        
     
    elif method == "shuffle_reduction":
        manipulated_text = shuffle_reduction(original_text)
    elif method == "change_topic":
        manipulated_text = change_topic(original_text, fake_dict)
    else:
        raise ValueError("Unknown manipulation method.")

    
    output_file="static/data/"+file
    write_docx(manipulated_text, output_file)
    sensitive_data_in_docx("static/data/"+file,"static/data/"+file)

    
    
    '''output_file = file_path.replace(ext, f"_{method}{ext}")
    if ext == '.docx':
        write_docx(manipulated_text, output_file)
    elif ext == '.pdf':
        write_pdf(manipulated_text, output_file)'''

    #print(f"File manipulated using '{method}' and saved as '{output_file}'.")


@app.route('/attack',methods=['POST','GET'])
def attack():
    msg=""
    st=""
    data=[]

    #ff=open("static/detect.txt","w")
    #ff.write("")
    #ff.close()

    ff=open("static/sms.txt","w")
    ff.write("1")
    ff.close()

            
    if request.method == "POST":
        query = request.form["query"]
        query_processed = preprocess_text(query)  # Preprocess query
        query_vector = vectorizer.transform([query_processed])  # Convert query to vector
        query_cluster = kmeans.predict(query_vector)[0]  # Predict cluster
        
        # Retrieve documents in the same cluster
        matched_docs = [filenames[i] for i in range(len(filenames)) if kmeans.labels_[i] == query_cluster]

        print(matched_docs)
        n_files=len(matched_docs)
        if n_files>0:
            path_main = 'static/data'

            ff=open("static/detect.txt","w")
            ff.write("1")
            ff.close()
            
            for fname in os.listdir(path_main):
                try:
                    if os.path.isfile("static/data/"+fname):
                        os.remove("static/data/"+fname)
                except:
                    print("no file")

            for fname1 in matched_docs:
                shutil.copy("static/ip_documents/"+fname1,"static/data/"+fname1)

                
                ext=fname1.split(".")
                if ext[1]=="docx":
                    print("yes")
                    manipulate_file(fname1,"basic_shuffle")
                if ext[1]=="pdf":

                    redaction_boxes = {
                        0: [(100, 500, 200, 550), (300, 400, 400, 450)],  # Page 0 redactions
                        1: [(150, 600, 250, 650)],                        # Page 1 redactions
                    }
                    data_pdf("static/data/"+fname1,"static/data/"+fname1,redaction_boxes)

                
            st="1"
        
        #return render_template("attack.html", query=query, matched_docs=matched_docs,st=st)
        # query=None, matched_docs=None

    return render_template("web/attack.html",msg=msg,st=st)

@app.route('/process1',methods=['POST','GET'])
def process1():
    msg=""
    dfile=[]
    sms_st=""
    act=request.args.get("act")
    val=""
    mess=""
    imgpath=""
    txtdata=""
    fnn=request.args.get("fnn")
    fs=request.args.get("fs")
    
    ff=open("static/det.txt","r")
    det=ff.read()
    ff.close()

    ff=open("static/sms.txt","r")
    sms=ff.read()
    ff.close()
    
    cc=det.split("|")
    email=cc[0]
    mobile=cc[1]
    
    ff=open("static/detect.txt","r")
    val=ff.read()
    ff.close()

    if val=="1":
        msg="attack"
        mess="Adversary Attack the Files"
        n=int(sms)
        if n==1:
            sms_st="1"
        
        
        n1=n+1
        sms1=str(n1)
        ff=open("static/sms.txt","w")
        ff.write(sms1)
        ff.close()
        
       
    elif val=="2":
        msg="attack"

    elif val=="3":
        msg="show"


    if val=="":
        msg="no"

    return render_template("web/process1.html", msg=msg,act=act,dfile=dfile,val=val,mess=mess,email=email,mobile=mobile,sms_st=sms_st)

@app.route('/attack1',methods=['POST','GET'])
def attack1():
    msg=""
    dfile=[]
    sms_st=""
    act=request.args.get("act")
    val=""
    mess=""
    imgpath=""
    txtdata=""
    fnn=request.args.get("fnn")
    fs=request.args.get("fs")
    
    ff=open("static/det.txt","r")
    det=ff.read()
    ff.close()

    ff=open("static/sms.txt","r")
    sms=ff.read()
    ff.close()
    
    cc=det.split("|")
    email=cc[0]
    mobile=cc[1]
    
    ff=open("static/detect.txt","r")
    val=ff.read()
    ff.close()


    msg="show"
    path_main="static/data"
    for fname in os.listdir(path_main):
        dt=[]
        fst=fname.split(".")
        dt.append(fst[1])
        dt.append(fname)
        dfile.append(dt)
  
    if act=="view":
        pt="static/data"
        
        if fs=="docx":
            fpath=pt
            txtdata=extract_text_from_docx(pt+"/"+fnn)
     
        elif fs=="pdf":
            fpath=pt
            imgpath=pt+"/"+fnn
        else:
            fpath=pt
            ff=open(pt+"/"+fnn,"r")
            txtdata=ff.read()
            ff.close()


    return render_template("web/attack1.html", msg=msg,act=act,dfile=dfile,val=val,mess=mess,email=email,mobile=mobile,sms_st=sms_st,fnn=fnn,fs=fs,imgpath=imgpath,txtdata=txtdata)

def highlight_differences(text1, text2):
    dmp = diff_match_patch()
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    diff_html = dmp.diff_prettyHtml(diffs)
    return diff_html

@app.route('/check',methods=['POST','GET'])
def check():
    msg=""
    dfile=[]
    sms_st=""
    mess=""
    act=request.args.get("act")
    text1=""
    text2=""
    differences=""
    ff=open("static/detect.txt","r")
    val=ff.read()
    ff.close()

    ff=open("static/sms.txt","r")
    sms=ff.read()
    ff.close()

    ff=open("static/det.txt","r")
    det=ff.read()
    ff.close()

    cc=det.split("|")
    email=cc[0]
    mobile=cc[1]

    if val=="1":
        msg="attack"
        ff=open("static/detect.txt","w")
        ff.write("2")
        ff.close()

        mess="Adversary Attack the Files"
        n=int(sms)
        if n==1:
            sms_st="1"
        
        
        n1=n+1
        sms1=str(n1)
        ff=open("static/sms.txt","w")
        ff.write(sms1)
        ff.close()

        
    elif val=="2":
        msg="dord"
        dfn=[]
        path_main="static/data"
        for fname in os.listdir(path_main):
            dfn.append(fname)
           
        fn1=dfn[0]
        file1="static/ip_documents/"+fn1
        shutil.copy("static/ip_documents/"+fn1,"static/upload/"+fn1)
        
        file2="static/upload/"+fn1

        #########
        original_text = read_docx("static/ip_documents/"+fn1)
        # Manipulation parameters
        decoy_words = ['quantum', 'blockchain', 'nanotech', 'space-time', 'encryption']
        fake_dict = {'data': 'watermelon', 'algorithm': 'sandwich', 'system': 'suitcase',
                     'network': 'cloud', 'model': 'pineapple'}

        if act=="1":
            manipulated_text = basic_shuffle(original_text)
            output_file="static/upload/"+fn1
            write_docx(manipulated_text, output_file)
            
            text1 = extract_text_from_docx(file1)
            text2 = extract_text_from_docx(file2)

            differences = highlight_differences(text1, text2)

        elif act=="2":
            manipulated_text = shuffle_increment(original_text, decoy_words)
            output_file="static/upload/"+fn1
            write_docx(manipulated_text, output_file)
            sensitive_data_in_docx("static/upload/"+fn1,"static/upload/"+fn1)
            
            text1 = extract_text_from_docx(file1)
            text2 = extract_text_from_docx(file2)

            differences = highlight_differences(text1, text2)
        elif act=="3":
            manipulated_text = shuffle_reduction(original_text)
            output_file="static/upload/"+fn1
            write_docx(manipulated_text, output_file)
            
            text1 = extract_text_from_docx(file1)
            text2 = extract_text_from_docx(file2)

            differences = highlight_differences(text1, text2)
        elif act=="4":
            manipulated_text = change_topic(original_text, fake_dict)
            output_file="static/upload/"+fn1
            write_docx(manipulated_text, output_file)
            
            text1 = extract_text_from_docx(file1)
            text2 = extract_text_from_docx(file2)

            differences = highlight_differences(text1, text2)

        elif act=="5":
            path_main="static/data"
            for fname in os.listdir(path_main):
                dt=[]
                fst=fname.split(".")
                dt.append(fst[1])
                dt.append(fname)
                dfile.append(dt)

            ff=open("static/detect.txt","w")
            ff.write("3")
            ff.close()
    elif val=="3":
        msg="show"
        path_main="static/data"
        for fname in os.listdir(path_main):
            dfile.append(fname)

    if val=="":
        msg="no"

    return render_template("check.html", msg=msg,act=act,dfile=dfile,differences=differences,text1=text1,mess=mess,email=email,mobile=mobile,sms_st=sms_st)



@app.route('/user_doc', methods=['GET', 'POST'])
def user_doc():
    uname=""
    msg=""
    act = request.args.get('act')
    if 'username' in session:
        uname = session['username']
    ff=open("user.txt",'r')
    uname=ff.read()
    ff.close()
    mycursor = mydb.cursor()

    mycursor.execute("SELECT * FROM ins_user where uname=%s",(uname,))
    value = mycursor.fetchone()
    
    mycursor.execute("SELECT * FROM ins_data")
    data = mycursor.fetchall()


    if act=="del":
        did = request.args.get('did')
        mycursor.execute("delete from ins_data where id=%s",(did,))
        mydb.commit()
        return redirect(url_for('user_doc')) 

    return render_template('user_doc.html',data=data)

@app.route('/down', methods=['GET', 'POST'])
def down():
    fn = request.args.get('fn')
    path = "static/ip_documents/"+fn
  
    return send_file(path, as_attachment=True)
@app.route('/down1', methods=['GET', 'POST'])
def down1():
    fn = request.args.get('fn')
    path = "static/data/"+fn
  
    return send_file(path, as_attachment=True)

@app.route('/logout')
def logout():
    # remove the username from the session if it is there
    session.pop('username', None)
    return redirect(url_for('index'))


if __name__ == "__main__":
    app.secret_key = os.urandom(12)
    app.run(debug=True,host='0.0.0.0', port=5000)
